from __future__ import annotations

import io
from typing import Optional

from logs.custom_logger import get_logger

logger = get_logger(__name__)

_NATIVE_MIMES = {"application/pdf"}
_NATIVE_EXTS = {"pdf", "png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff", "tif"}

_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOCX_EXTS = {"docx"}

_RTF_MIMES = {"application/rtf", "text/rtf"}
_RTF_EXTS = {"rtf"}

_TEXT_MIMES = {
    "application/json",
    "application/xml",
    "application/x-yaml",
    "application/yaml",
    "application/csv",
}
_TEXT_EXTS = {
    "txt", "csv", "tsv", "md", "markdown", "json", "log",
    "xml", "html", "htm", "yaml", "yml", "ini", "rst",
}


def _ext_of(file_name: Optional[str]) -> str:
    if file_name and "." in file_name:
        return file_name.rsplit(".", 1)[1].lower().strip()
    return ""


def _norm_mime(media_type: Optional[str]) -> str:
    return (media_type or "").split(";", 1)[0].strip().lower()


def _classify(media_type: Optional[str], file_name: Optional[str]) -> Optional[str]:
    ext = _ext_of(file_name)
    mime = _norm_mime(media_type)

    if ext in _NATIVE_EXTS or mime in _NATIVE_MIMES or mime.startswith("image/"):
        return "native"
    if ext in _DOCX_EXTS or mime in _DOCX_MIMES:
        return "docx"
    if ext in _RTF_EXTS or mime in _RTF_MIMES:
        return "rtf"
    if ext in _TEXT_EXTS or mime in _TEXT_MIMES or mime.startswith("text/"):
        return "text"
    return None


def _decode_text(raw: bytes) -> str:
    """Decode bytes to str, detecting the encoding when possible."""
    if not raw:
        return ""
    try:
        from charset_normalizer import from_bytes  

        best = from_bytes(raw).best()
        if best is not None:
            return str(best)
    except Exception:
        pass
    return raw.decode("utf-8", errors="replace")


def _extract_docx(raw: bytes) -> str:
    import docx  

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_rtf(raw: bytes) -> str:
    from striprtf.striprtf import rtf_to_text  

    return rtf_to_text(_decode_text(raw))


def extract_text(
    raw: bytes,
    media_type: Optional[str] = None,
    file_name: Optional[str] = None,
) -> Optional[str]:
    kind = _classify(media_type, file_name)

    if kind == "native":
        return None
    if kind is None:
        logger.info(
            "text_extraction: unsupported file type, skipping",
            media_type=media_type,
            file_name=file_name,
        )
        return None

    try:
        if kind == "docx":
            text = _extract_docx(raw)
        elif kind == "rtf":
            text = _extract_rtf(raw)
        else:  # "text"
            text = _decode_text(raw)
    except Exception as e:
        logger.warning(
            "text_extraction: failed to extract text",
            kind=kind,
            media_type=media_type,
            file_name=file_name,
            error=str(e),
            severity="medium",
        )
        return None

    text = (text or "").strip()
    if not text:
        return None
    return text
